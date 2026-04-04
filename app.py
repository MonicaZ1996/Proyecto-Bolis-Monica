from flask import Flask, render_template, request, redirect, send_file
from conexion.conexion import get_connection

# PDF bonito
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

# WORD
from docx import Document

app = Flask(__name__)

# ================= INICIO =================
@app.route("/")
def inicio():
    return render_template("index.html")


# ================= PRODUCTOS =================
@app.route("/producto")
def producto():
    producto = db.connection.cursor()
    producto.execute("SELECT * FROM producto")
    datos = producto.fetchall()
    return render_template("producto.html", producto=datos)


# ================= CREAR =================
@app.route("/producto/crear", methods=["GET", "POST"])
def crear():
    if request.method == "POST":
        try:
            nombre = request.form["nombre"]
            precio = request.form["precio"]
            stock = request.form["stock"]

            conn = get_connection()
            cursor = conn.cursor()

            cursor.execute(
                "INSERT INTO producto (nombre, precio, stock) VALUES (%s,%s,%s)",
                (nombre, precio, stock)
            )

            conn.commit()
            conn.close()
        except Exception as e:
            print("Error:", e)

        return redirect("/producto")

    return render_template("producto/crear.html")


# ================= EDITAR =================
@app.route("/producto/editar/<int:id>", methods=["GET", "POST"])
def editar(id):
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)

        if request.method == "POST":
            nombre = request.form["nombre"]
            precio = request.form["precio"]
            stock = request.form["stock"]

            cursor.execute(
                "UPDATE producto SET nombre=%s, precio=%s, stock=%s WHERE id_producto=%s",
                (nombre, precio, stock, id)
            )

            conn.commit()
            conn.close()
            return redirect("/producto")

        cursor.execute("SELECT * FROM producto WHERE id_producto=%s", (id,))
        producto = cursor.fetchone()

        conn.close()

    except Exception as e:
        print("Error:", e)
        producto = None

    return render_template("producto/editar.html", producto=producto)


# ================= ELIMINAR =================
@app.route("/producto/eliminar/<int:id>")
def eliminar(id):
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("DELETE FROM producto WHERE id_producto=%s", (id,))
        conn.commit()

        conn.close()
    except Exception as e:
        print("Error:", e)

    return redirect("/producto")


# ================= CONTACTO =================
@app.route("/contacto")
def contacto():
    return render_template("contacto.html")


# ================= DATOS =================
@app.route("/datos")
def datos():
    return render_template("datos.html")


# ================= PDF =================
@app.route("/reporte")
def reporte():
    archivo = "reporte.pdf"
    doc = SimpleDocTemplate(archivo)

    elementos = []
    estilos = getSampleStyleSheet()

    elementos.append(Paragraph("Informe de Productos - Bolis", estilos["Title"]))

    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT nombre, precio, stock FROM productos")
        datos = cursor.fetchall()
        conn.close()
    except:
        datos = []

    data = [["Nombre", "Precio", "Stock"]]
    for p in datos:
        data.append([p[0], str(p[1]), str(p[2])])

    tabla = Table(data)
    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.pink),
        ("GRID", (0,0), (-1,-1), 1, colors.black)
    ]))

    elementos.append(tabla)
    doc.build(elementos)

    return send_file(archivo, as_attachment=True)


# ================= WORD =================
@app.route("/reporte_word")
def reporte_word():
    doc = Document()
    doc.add_heading("Informe de Productos - Bolis", 0)

    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT nombre, precio, stock FROM productos")
        datos = cursor.fetchall()
        conn.close()
    except:
        datos = []

    for p in datos:
        doc.add_paragraph(f"Nombre: {p[0]} | Precio: {p[1]} | Stock: {p[2]}")

    archivo = "reporte.docx"
    doc.save(archivo)

    return send_file(archivo, as_attachment=True)


# ================= RUN =================
if __name__ == "__main__":
    app.run(debug=True)

# listar productos
@app.route("/producto")
def producto():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT * FROM productos")
    datos = cursor.fetchall()

    conn.close()

    return render_template("producto.html", producto=datos)

# insertar producto
@app.route("/producto/crear", methods=["GET", "POST"])
def crear_producto():
    if request.method == "POST":
        nombre = request.form["nombre"]
        precio = request.form["precio"]
        stock = request.form["stock"]

        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute(
            "INSERT INTO productos (nombre, precio, stock) VALUES (%s,%s,%s)",
            (nombre, precio, stock)
        )

        conn.commit()
        conn.close()

        return redirect("/producto")

    return render_template("crear.html")
